from docx import Document
import logging

logger = logging.getLogger(__name__)


def process_docx(file_path: str) -> dict:
    """Extract text from a DOCX file including paragraphs and tables."""
    try:
        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        text = "\n".join(text_parts).strip()
        if not text:
            raise ValueError("No text could be extracted from the DOCX file")

        logger.info(f"Extracted {len(text)} characters from DOCX: {file_path}")
        return {"text": text, "type": "docx"}

    except Exception as e:
        logger.error(f"Failed to process DOCX {file_path}: {e}")
        raise
